"""Unit tests for the format renderers (feature 8, architecture.md §8).

These use a real ``tmp_path`` directory (no filesystem mocking, per
conventions §7) and assert concrete results: the ``.docx`` is reopened with
``python-docx`` and its content checked against the input Markdown.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from thesis_agents.adapters.formats import render_to_format, slugify
from thesis_agents.schemas.models import DocSpec


def _make_spec(fmt: str, title: str = "My Thesis Chapter") -> DocSpec:
    return DocSpec(
        title=title,
        docType="thesis-chapter",
        format=fmt,
        language="en",
        chapter={"number": 1, "title": "Introduction"},
        audience="committee",
        targetWords=2000,
        citationStyle="APA",
        requirements=["cover the background"],
        notes="",
    )


MARKDOWN = """# Main Heading

This is an intro paragraph with **bold text** and *italic text*.

## Subsection

- first bullet
- second bullet

1. first numbered
2. second numbered
"""


def test_render_to_format_docx_writes_file_reflecting_markdown(tmp_path: Path) -> None:
    spec = _make_spec("docx")

    result = render_to_format(MARKDOWN, spec, tmp_path)

    assert result.exists()
    assert result.suffix == ".docx"
    assert result.parent == tmp_path

    document = Document(str(result))
    texts = [p.text for p in document.paragraphs]

    # Heading text from the input Markdown is present.
    assert "Main Heading" in texts
    assert "Subsection" in texts
    # A paragraph reflecting the input body (bold/italic markers stripped).
    assert "This is an intro paragraph with bold text and italic text." in texts
    # Bullet and numbered item text survives.
    assert "first bullet" in texts
    assert "first numbered" in texts


def test_render_to_format_docx_marks_bold_run(tmp_path: Path) -> None:
    spec = _make_spec("docx")

    result = render_to_format("Body with **strongword** inside.", spec, tmp_path)

    document = Document(str(result))
    bold_runs = [
        run.text
        for paragraph in document.paragraphs
        for run in paragraph.runs
        if run.bold
    ]
    italic_source = render_to_format(
        "Body with *leaning* inside.", _make_spec("docx", "Other"), tmp_path
    )
    italic_doc = Document(str(italic_source))
    italic_runs = [
        run.text
        for paragraph in italic_doc.paragraphs
        for run in paragraph.runs
        if run.italic
    ]

    assert "strongword" in bold_runs
    assert "leaning" in italic_runs


def test_render_to_format_uses_name_hint_for_filename(tmp_path: Path) -> None:
    spec = _make_spec("docx", title="Ignored Title")

    result = render_to_format("# X", spec, tmp_path, name_hint="Custom Name!")

    assert result.name == "custom-name.docx"


def test_render_to_format_creates_missing_output_dir(tmp_path: Path) -> None:
    spec = _make_spec("docx")
    target = tmp_path / "nested" / "out"

    result = render_to_format("# X", spec, target)

    assert result.exists()
    assert result.parent == target


def test_render_to_format_pptx_raises_not_implemented(tmp_path: Path) -> None:
    spec = _make_spec("pptx")

    with pytest.raises(NotImplementedError):
        render_to_format("# X", spec, tmp_path)


@pytest.mark.parametrize(
    ("raw", "expected"),
    [
        ("My Thesis Chapter", "my-thesis-chapter"),
        ("  Trailing / Spaces  ", "trailing-spaces"),
        ("Chapter 1: Intro!!! (draft)", "chapter-1-intro-draft"),
        ("Hello___World", "hello-world"),
        ("MiXeD CaSe", "mixed-case"),
        ("", "document"),
        ("***", "document"),
        ("---already---", "already"),
    ],
)
def test_slugify_produces_safe_slug(raw: str, expected: str) -> None:
    assert slugify(raw) == expected
