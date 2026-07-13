"""Markdown → ``.docx`` renderer (architecture.md §8).

Renders the *Writer's Markdown subset* — deliberately small, not full
CommonMark — to a Word document via ``python-docx`` (imported as ``docx``):

- ``#`` / ``##`` / ``###`` headings → Word heading levels 1/2/3;
- blank-line-separated paragraphs;
- ``-`` / ``*`` bullet list items → ``List Bullet`` paragraphs;
- ``1.`` numbered list items → ``List Number`` paragraphs;
- inline ``**bold**`` and ``*italic*`` spans → bold / italic runs.

Anything else is treated as plain paragraph text. The renderer takes the target
``output_path`` from its caller (the controller passes a path under the
config-provided output dir); it reads no environment and touches nothing else.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from thesis_agents.schemas.models import DocSpec

#: A heading line: one to three ``#`` then whitespace then the heading text.
_HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
#: A bullet list item: ``-`` or ``*`` then whitespace then the item text.
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
#: A numbered list item: digits, a dot, whitespace, then the item text.
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)$")
#: Inline emphasis spans: ``**bold**`` (checked first) or ``*italic*``.
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def _add_inline(paragraph, text: str) -> None:
    """Add ``text`` to ``paragraph`` as runs, honoring ``**bold**``/``*italic*``."""
    for token in _INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def render_docx(markdown: str, spec: DocSpec, output_path: Path) -> Path:
    """Render the Markdown subset to ``output_path`` as a ``.docx`` file.

    :param markdown: the Writer's Markdown (subset documented in the module).
    :param spec: the document spec (unused by the renderer beyond routing, kept
        for signature parity with :func:`slides.render_pptx`).
    :param output_path: absolute path of the ``.docx`` file to write.
    :returns: ``output_path`` after the document has been saved.
    """
    document = Document()

    for line in markdown.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        heading = _HEADING_RE.match(stripped)
        if heading is not None:
            level = len(heading.group(1))
            document.add_heading(heading.group(2).strip(), level=level)
            continue

        bullet = _BULLET_RE.match(stripped)
        if bullet is not None:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline(paragraph, bullet.group(1).strip())
            continue

        numbered = _NUMBERED_RE.match(stripped)
        if numbered is not None:
            paragraph = document.add_paragraph(style="List Number")
            _add_inline(paragraph, numbered.group(1).strip())
            continue

        paragraph = document.add_paragraph()
        _add_inline(paragraph, stripped)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path
